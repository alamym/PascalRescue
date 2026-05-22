import os
import re
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def parse_wre(txt_path):
    reports = {}
    current_id = None

    if not os.path.exists(txt_path):
        return reports

    with open(txt_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    for line in lines:
        line = line.strip()
        if not line:
            continue

        # 過濾雜訊
        if re.match(r'^\d+ of \d+$', line): continue
        if 'REPORT ON THE EXAMINATION' in line: continue

        # 偵測大題標題 - 如果遇到新的 Question，終止當前小題抓取
        if re.match(r'^\d+\tQuestion', line):
            current_id = None
            continue

        # 偵測小題編號 (e.g., 01.1)
        id_match = re.match(r'^(\d+\.\d+)\s+(.*)', line)
        if id_match:
            current_id = id_match.group(1)
            content = id_match.group(2)
            if current_id not in reports:
                reports[current_id] = []
            reports[current_id].append(content)
        elif current_id:
            # 後續段落
            reports[current_id].append(line)

    return reports

def integrate_wre_to_docx(docx_input, wre_txt, docx_output):
    print(f"Reading WRE: {wre_txt}")
    wre_data = parse_wre(wre_txt)

    print(f"Opening MS Docx: {docx_input}")
    doc = Document(docx_input)

    # 建立輸出目錄
    os.makedirs(os.path.dirname(docx_output), exist_ok=True)

    new_doc = Document()
    # 複製原始文件的節點屬性
    new_doc.sections[0].page_height = doc.sections[0].page_height
    new_doc.sections[0].page_width = doc.sections[0].page_width
    new_doc.sections[0].left_margin = doc.sections[0].left_margin
    new_doc.sections[0].right_margin = doc.sections[0].right_margin
    new_doc.sections[0].top_margin = doc.sections[0].top_margin
    new_doc.sections[0].bottom_margin = doc.sections[0].bottom_margin

    for para in doc.paragraphs:
        text = para.text.strip()

        # 1. 處理大題標題：加上 2019-
        # 匹配 Question 1, Question 2 等
        if re.match(r'^Question\s+\d+', text, re.I) and '.' not in text:
            new_para = new_doc.add_paragraph()
            run = new_para.add_run(f"2019-{text}")
            run.bold = True
            run.font.name = 'Calisto MT'
            run.font.size = Pt(14)
            continue

        # 2. 正常加入原段落
        new_para = new_doc.add_paragraph()
        for run in para.runs:
            new_run = new_para.add_run(run.text)
            new_run.bold = run.bold
            new_run.italic = run.italic
            new_run.font.name = run.font.name
            new_run.font.size = run.font.size

        new_para.paragraph_format.left_indent = para.paragraph_format.left_indent
        new_para.paragraph_format.first_line_indent = para.paragraph_format.first_line_indent
        new_para.alignment = para.alignment

        # 3. 偵測題號並插入 WRE
        id_match = re.match(r'^(\d+\.\d+)', text)
        if id_match:
            qid = id_match.group(1)
            if qid in wre_data:
                new_doc.add_paragraph("")

                header_para = new_doc.add_paragraph()
                header_run = header_para.add_run("Examiners’ reports on individual questions")
                header_run.bold = True
                header_run.font.name = 'Calisto MT'
                header_run.font.size = Pt(12)

                for report_text in wre_data[qid]:
                    report_para = new_doc.add_paragraph()
                    r_run = report_para.add_run(report_text)
                    r_run.font.name = 'Calisto MT'
                    r_run.font.size = Pt(12)

                new_doc.add_paragraph("")

    print(f"Saving result to: {docx_output}")
    new_doc.save(docx_output)
    print("Integration successful!")

# 執行 2019 整合
integrate_wre_to_docx(
    "data/outputs/docx/AQA-8464-B-1F-MS-19.docx",
    "data/txts/WRE/AQA-8464-B-1F-WRE-19.txt",
    "data/outputs/WRE/AQA-8464-B-1F-MS-19.docx"
)
